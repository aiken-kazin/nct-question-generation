#!/usr/bin/env python3
"""Build one .docx per model for MATH from the per-question files (so figure
paths are available). Standalone grouped by level (image questions embed the
figure); then context blocks."""
import json
from collections import defaultdict
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = Path("output/math_final/math")
OUTDIR = Path("output/math_final/docx")
OUTDIR.mkdir(parents=True, exist_ok=True)
MODELS = ["google/gemini-3.1-pro-preview", "openai/gpt-5.5", "anthropic/claude-sonnet-4.6"]


def slug(m): return m.split("/")[-1]


def run(p, t, bold=False, italic=False, size=11, color=None):
    r = p.add_run(t); r.bold = bold; r.italic = italic; r.font.size = Pt(size)
    if color: r.font.color.rgb = RGBColor(*color)
    return r


def q_block(doc, n, d, level=None, topic=None, fmt=None, figure_path=None):
    p = doc.add_paragraph()
    run(p, f"{n}. ", bold=True)
    tags = [t for t in [f"Деңгей {level}" if level else None, topic,
                        "сурет" if fmt == "image" else None] if t]
    if tags:
        run(p, "[" + " | ".join(tags) + "]  ", bold=True, color=(0x55, 0x55, 0x55), size=9)
    run(p, d["question_text"])
    if fmt == "image" and figure_path and Path(figure_path).is_file():
        fp = doc.add_paragraph(); fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        try:
            fp.add_run().add_picture(figure_path, width=Inches(3.2))
        except Exception:
            run(doc.add_paragraph(), f"[сурет: {figure_path}]", italic=True, size=9)
    for o in d["options"]:
        op = doc.add_paragraph(style="List Bullet")
        ok = o["label"] == d["correct_answer"]
        run(op, f"{o['label']}) {o['text']}", bold=ok)
        if ok: run(op, "  ✓", bold=True, color=(0x1B, 0x7F, 0x3A))
    pa = doc.add_paragraph()
    run(pa, "Жауабы: ", bold=True); run(pa, d["correct_answer"], bold=True, color=(0x1B, 0x7F, 0x3A))
    sc = d.get("critic_score")
    if sc is not None:
        run(pa, f"    (критик: {sc:.1f}/10)", italic=True, color=(0x77, 0x77, 0x77), size=9)
    ver = (d.get("critic_feedback") or {}).get("verification")
    if ver and ver.get("ran"):
        ok = ver.get("passed")
        run(pa, f"   [SymPy: {'расходится' if ver.get('contradicted') else 'подтверждено' if ok else 'проверено'}]",
            italic=True, size=9, color=(0x1B, 0x7F, 0x3A) if ok else (0xB0, 0x30, 0x30))
    if d.get("explanation"):
        pe = doc.add_paragraph()
        run(pe, "Түсіндірме: ", bold=True, size=9, color=(0x55, 0x55, 0x55))
        run(pe, d["explanation"], italic=True, size=9, color=(0x55, 0x55, 0x55))
    doc.add_paragraph()


for model in MODELS:
    mdir = ROOT / slug(model)
    if not mdir.is_dir():
        continue
    standalone = []
    for f in sorted(mdir.glob("level_*/*.json")):
        if f.name.startswith("_"): continue
        standalone.append(json.loads(f.read_text(encoding="utf-8")))
    blocks = [json.loads(f.read_text(encoding="utf-8")) for f in sorted((mdir / "context").glob("*.json"))]

    doc = Document()
    doc.add_heading("НТЦ — Математика: генерленген тест тапсырмалары", level=0)
    run(doc.add_paragraph(), f"Модель: {model}", bold=True, size=12)
    nq_ctx = sum(len(b.get("questions", [])) for b in blocks)
    run(doc.add_paragraph(),
        f"Standalone: {len(standalone)} | Контекстік блоктар: {len(blocks)} ({nq_ctx} сұрақ) | Барлығы: {len(standalone)+nq_ctx}",
        italic=True, color=(0x55, 0x55, 0x55), size=10)

    doc.add_heading("I. Контекстсіз тапсырмалар (standalone)", level=1)
    by_lvl = defaultdict(list)
    for d in standalone:
        by_lvl[d.get("level", "?")].append(d)
    n = 1
    for lvl in ["A", "B", "C"]:
        items = by_lvl.get(lvl, [])
        if not items: continue
        doc.add_heading(f"Деңгей {lvl} ({len(items)})", level=2)
        for d in items:
            q_block(doc, n, d, level=d.get("level"), topic=d.get("topic"),
                    fmt=d.get("format"), figure_path=d.get("figure_path"))
            n += 1

    if blocks:
        doc.add_heading("II. Мәнмәтіндік тапсырмалар (context blocks)", level=1)
        for bi, b in enumerate(blocks, 1):
            doc.add_heading(f"Блок {bi}", level=2)
            pp = doc.add_paragraph(); run(pp, "Мәтін: ", bold=True); run(pp, b.get("passage", ""))
            doc.add_paragraph()
            for qi, q in enumerate(b.get("questions", []), 1):
                q_block(doc, qi, q)

    out = OUTDIR / f"Math_{slug(model)}.docx"
    doc.save(out)
    print(f"wrote {out}  (standalone={len(standalone)}, context_q={nq_ctx})")
