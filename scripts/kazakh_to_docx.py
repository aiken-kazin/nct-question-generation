#!/usr/bin/env python3
"""Build one .docx per model from all_kazakh_questions.json:
questions + options (correct marked) + answer + topic + level + explanation +
ensemble score. No-context grouped by level; then context blocks."""
import json
from collections import defaultdict
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

SRC = Path("output/kazakh_final/all_kazakh_questions.json")
OUTDIR = Path("output/kazakh_final/docx")
OUTDIR.mkdir(parents=True, exist_ok=True)

data = json.load(open(SRC, encoding="utf-8"))
MODELS = ["google/gemini-3.1-pro-preview", "openai/gpt-5.5", "anthropic/claude-sonnet-4.6"]


def slug(m):
    return m.split("/")[-1]


def run(p, text, bold=False, italic=False, size=11, color=None):
    r = p.add_run(text); r.bold = bold; r.italic = italic; r.font.size = Pt(size)
    if color:
        r.font.color.rgb = RGBColor(*color)
    return r


def add_question(doc, n, q, level=None, topic=None):
    p = doc.add_paragraph()
    head = f"{n}. "
    tags = []
    if level:
        tags.append(f"Деңгей {level}")
    if topic:
        tags.append(topic)
    run(p, head, bold=True)
    if tags:
        run(p, "[" + " | ".join(tags) + "]  ", bold=True, color=(0x55, 0x55, 0x55), size=9)
    run(p, q["question_text"])
    for o in q["options"]:
        op = doc.add_paragraph(style="List Bullet")
        correct = o["label"] == q["correct_answer"]
        run(op, f"{o['label']}) {o['text']}", bold=correct)
        if correct:
            run(op, "  ✓", bold=True, color=(0x1B, 0x7F, 0x3A))
    pa = doc.add_paragraph()
    run(pa, "Жауабы: ", bold=True); run(pa, q["correct_answer"], bold=True, color=(0x1B, 0x7F, 0x3A))
    sc = q.get("critic_score")
    if sc is not None:
        run(pa, f"    (критик: {sc:.1f}/10)", italic=True, color=(0x77, 0x77, 0x77), size=9)
    if q.get("explanation"):
        pe = doc.add_paragraph()
        run(pe, "Түсіндірме: ", bold=True, size=9, color=(0x55, 0x55, 0x55))
        run(pe, q["explanation"], italic=True, size=9, color=(0x55, 0x55, 0x55))
    doc.add_paragraph()


for model in MODELS:
    nc = [q for q in data["no_context"] if q["model"] == model]
    blocks = [b for b in data["with_context"] if b["model"] == model]
    if not nc and not blocks:
        continue
    doc = Document()
    doc.add_heading(f"НТЦ — Қазақ тілі: генерленген тест тапсырмалары", level=0)
    sub = doc.add_paragraph()
    run(sub, f"Модель: {model}", bold=True, size=12)
    cnt = doc.add_paragraph()
    nq_ctx = sum(len(b["questions"]) for b in blocks)
    run(cnt, f"Standalone: {len(nc)} | Контекстік блоктар: {len(blocks)} ({nq_ctx} сұрақ) | Барлығы: {len(nc)+nq_ctx}",
        italic=True, color=(0x55, 0x55, 0x55), size=10)

    # No-context grouped by level
    doc.add_heading("I. Контекстсіз тапсырмалар (standalone)", level=1)
    by_lvl = defaultdict(list)
    for q in nc:
        by_lvl[q.get("level", "?")].append(q)
    n = 1
    for lvl in ["A", "B", "C"]:
        items = by_lvl.get(lvl, [])
        if not items:
            continue
        doc.add_heading(f"Деңгей {lvl} ({len(items)})", level=2)
        for q in items:
            add_question(doc, n, q, level=q.get("level"), topic=q.get("topic"))
            n += 1

    # Context blocks
    if blocks:
        doc.add_heading("II. Мәнмәтіндік тапсырмалар (context blocks)", level=1)
        for bi, b in enumerate(blocks, 1):
            doc.add_heading(f"Блок {bi}", level=2)
            pp = doc.add_paragraph()
            run(pp, "Мәтін: ", bold=True)
            run(pp, b.get("passage", ""))
            doc.add_paragraph()
            for qi, q in enumerate(b["questions"], 1):
                add_question(doc, qi, q)

    out = OUTDIR / f"Kazakh_{slug(model)}.docx"
    doc.save(out)
    print(f"wrote {out}  (standalone={len(nc)}, context_q={nq_ctx})")
